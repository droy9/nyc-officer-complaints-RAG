"""
Document Processing
====================
Handles extraction of text from various file formats.
"""

import logging
import uuid
from pathlib import Path
from typing import Optional, Dict, Any
from datetime import datetime

from .text_processor import TextProcessor, ProcessedDocument

logger = logging.getLogger("rag.document_processor")


class DocumentProcessor:
    """Processes uploaded documents (PDF, DOCX, TXT) into chunks."""
    
    def __init__(self, text_processor: TextProcessor):
        self.text_processor = text_processor
        self._pdf_available = False
        self._docx_available = False
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which document processing libraries are available."""
        try:
            import pymupdf
            self._pdf_available = True
            logger.info("✓ PDF processing available (PyMuPDF)")
        except ImportError:
            logger.warning("PyMuPDF not installed. PDF processing disabled. Install with: pip install pymupdf")
        
        try:
            import docx
            self._docx_available = True
            logger.info("✓ DOCX processing available (python-docx)")
        except ImportError:
            logger.warning("python-docx not installed. DOCX processing disabled. Install with: pip install python-docx")
    
    def process_file(
        self, 
        file_path: Path, 
        original_filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """
        Process a file and return a ProcessedDocument.
        
        Args:
            file_path: Path to the uploaded file
            original_filename: Original name of the uploaded file
            metadata: Optional additional metadata
            
        Returns:
            ProcessedDocument ready for indexing
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        # Extract text based on file type
        if suffix == '.pdf':
            text = self._extract_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            text = self._extract_docx(file_path)
        elif suffix == '.txt':
            text = self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        
        if not text or not text.strip():
            raise ValueError(f"No text content extracted from {original_filename}")
        
        # Chunk the text
        chunks = self.text_processor.chunk_text(text)
        
        # Generate document ID
        doc_id = str(uuid.uuid4())[:8]
        
        # Build metadata
        doc_metadata = {
            "uploaded_at": datetime.utcnow().isoformat(),
            "file_type": suffix[1:],  # Remove the dot
            "file_size": file_path.stat().st_size,
            "char_count": len(text),
            **(metadata or {})
        }
        
        logger.info(
            f"Processed {original_filename}: {len(text)} chars → {len(chunks)} chunks"
        )
        
        return ProcessedDocument(
            source_type="uploaded_document",
            document_id=doc_id,
            filename=original_filename,
            chunks=chunks,
            metadata=doc_metadata
        )
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if not self._pdf_available:
            raise ImportError("PyMuPDF not installed. Install with: pip install pymupdf")
        
        import pymupdf
        
        text_parts = []
        with pymupdf.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text())
        
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not self._docx_available:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        import docx
        
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @property
    def supported_types(self) -> list:
        """Return list of supported file types."""
        types = ['.txt']
        if self._pdf_available:
            types.append('.pdf')
        if self._docx_available:
            types.extend(['.docx', '.doc'])
        return types

